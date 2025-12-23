
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime

class DocxReport:
    def __init__(self):
        self.doc = Document()
        
    def generate(self, results, target_url, output_path, network_info=None):
        # Title
        self.doc.add_heading('Relatório Técnico de Segurança Web', 0)
        p = self.doc.add_paragraph('Pentest Automatizado - Análise de Vulnerabilidades')
        p.add_run(f"\nAlvo: {target_url}")
        p.add_run(f"\nData: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Network Info
        if network_info:
            self.doc.add_heading('Informações de Rede', level=1)
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'IP'
            hdr_cells[1].text = 'Latência'
            hdr_cells[2].text = 'DNS'
            hdr_cells[3].text = 'Server'
            
            row_cells = table.add_row().cells
            row_cells[0].text = network_info.get('ip', 'N/A')
            row_cells[1].text = network_info.get('latency_ms', 'N/A')
            row_cells[2].text = network_info.get('dns_time_ms', 'N/A')
            row_cells[3].text = network_info.get('server', 'N/A')

        # Findings
        self.doc.add_heading('Resultados Detalhados', level=1)
        
        for res in results:
            vuln = res.get('vulnerability', 'Unknown')
            severity = res.get('severity', 'Info')
            
            h = self.doc.add_heading(f"[{severity.upper()}] {vuln}", level=2)
            
            # Simple content
            self.doc.add_paragraph(f"Categoria: {res.get('category')}")
            self.doc.add_paragraph(f"Detalhes: {res.get('details')}")
            
            if res.get('manual_test') and res['manual_test'] != '-':
                self.doc.add_paragraph("Teste Manual:", style='Intense Quote')
                self.doc.add_paragraph(res['manual_test'])

        self.doc.save(output_path)
        return output_path

def generate_docx_report(results, target, path, net_info):
    try:
        reporter = DocxReport()
        return reporter.generate(results, target, path, net_info)
    except Exception as e:
        print(f"DOCX Generation Error: {e}")
        return None
